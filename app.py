import os
import io
import time
import re
import sqlite3
import datetime
import streamlit as st
import pandas as pd
import plotly.express as px
from google import genai
from google.genai import types
from PIL import Image
from pypdf import PdfReader
import docx
from gtts import gTTS

st.set_page_config(page_title="المنظومة الاستشارية والتنفيذية الشاملة", layout="wide", page_icon="🏛️")

# --- دالة تنظيف التوقيتات الصوتية (Timestamps) ---
def clean_text_output(text: str) -> str:
    if not text:
        return ""
    cleaned = re.sub(r'\b\d{1,2}:\d{2}\b', '', text)
    cleaned = re.sub(r' +', ' ', cleaned)
    return cleaned.strip()

# --- قاعدة البيانات المحلية الدائمة (SQLite) ---
DB_FILE = "archive.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            category TEXT,
            title TEXT,
            prompt TEXT,
            response TEXT
        )
    """)
    conn.commit()
    conn.close()

def save_record(category, title, prompt, response):
    try:
        conn = sqlite3.connect(DB_FILE)
        c = conn.cursor()
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        c.execute("INSERT INTO records (timestamp, category, title, prompt, response) VALUES (?, ?, ?, ?, ?)",
                  (ts, category, title, prompt, response))
        conn.commit()
        conn.close()
    except Exception as e:
        st.error(f"خطأ في حفظ الأرشيف: {e}")

def get_records(search_query="", category_filter="الكل"):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    query = "SELECT id, timestamp, category, title, prompt, response FROM records WHERE 1=1"
    params = []
    
    if category_filter != "الكل":
        query += " AND category = ?"
        params.append(category_filter)
        
    if search_query.strip():
        query += " AND (title LIKE ? OR prompt LIKE ? OR response LIKE ?)"
        s = f"%{search_query.strip()}%"
        params.extend([s, s, s])
        
    query += " ORDER BY id DESC"
    c.execute(query, params)
    rows = c.fetchall()
    conn.close()
    return rows

def delete_record(record_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM records WHERE id = ?", (record_id,))
    conn.commit()
    conn.close()

init_db()

# --- إعداد اتصال الذكاء الاصطناعي مع التبديل التلقائي عند الحصة 429 ---
api_key = st.secrets.get("GEMINI_API_KEY") or os.environ.get("GEMINI_API_KEY")
if not api_key:
    st.error("⚠️ يرجى ضبط مفتاح GEMINI_API_KEY في إعدادات Secrets.")
    st.stop()

client = genai.Client(api_key=api_key.strip())
MODELS = ["gemini-3.6-flash", "gemini-2.5-flash", "gemini-2.0-flash"]

def generate_with_retry(contents, max_retries=2):
    for model_name in MODELS:
        for attempt in range(max_retries):
            try:
                return client.models.generate_content(model=model_name, contents=contents)
            except Exception as e:
                err_str = str(e)
                if any(x in err_str for x in ["429", "503", "RESOURCE_EXHAUSTED", "UNAVAILABLE"]):
                    if attempt < max_retries - 1:
                        time.sleep(4)
                        continue
                break
    raise Exception("خوادم الذكاء الاصطناعي تشهد ضغطاً مؤقتاً، يرجى الانتظار 30 ثانية والمحاولة مجدداً.")

def text_to_audio_bytes(text_arabic):
    try:
        clean_text = text_arabic.replace("*", "").replace("#", "")[:500]
        tts = gTTS(text=clean_text, lang='ar', slow=False)
        fp = io.BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return fp.getvalue()
    except Exception:
        return None

def extract_text_from_file(file):
    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "\n".join([p.extract_text() or "" for p in reader.pages])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def create_docx_download(content_text, title="المستند"):
    doc = docx.Document()
    doc.add_heading(title, level=0)
    for paragraph in content_text.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def prepare_multimodal_payload(system_instruction, user_text, uploaded_file):
    """دالة شاملة لتجهيز النصوص والملفات والصور معاً لأي قسم"""
    payload = []
    file_info = ""
    
    if uploaded_file is not None:
        if uploaded_file.type.startswith("image/"):
            img = Image.open(uploaded_file)
            payload.append(img)
            file_info = f"\n[مرفق صورة مفحوصة: {uploaded_file.name}]"
        else:
            extracted = extract_text_from_file(uploaded_file)
            file_info = f"\n[محتوى المستند المرفق ({uploaded_file.name})]:\n{extracted[:12000]}\n"
            
    full_prompt = f"{system_instruction}\n{file_info}\n\n[المدخلات والملاحظات المطلوبة]:\n{user_text}\n"
    payload.append(full_prompt)
    return payload

# --- الشريط الجانبي ---
with st.sidebar:
    st.header("👤 خيارات الوكيل المساعد")
    secretary_mode = st.selectbox(
        "نبرة الصياغة:",
        [
            "مساعد وسكرتير تنفيذي شامل ورصين",
            "مستشار فكري وحكيم ناصح",
            "قائد استراتيجي وإداري صارم",
            "باحث ومحقق علمي دقيق",
            "خبير تصاميم وديكورات",
            "مستشار زراعي وصحي"
        ]
    )
    voice_output = st.checkbox("🔊 نطق الردود صوتياً", value=True)
    st.markdown("---")
    st.caption("✅ جميع الأقسام تدعم رفع الملفات (PDF / Word) والصور والنصوص وتحفظ في الأرشيف الدائم.")

st.title("🏛️ المنظومة الاستشارية والتنفيذية الشاملة")

tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab_arch = st.tabs([
    "🎙️ السكرتير الصوتي والشخصي",
    "📚 البحوث الحوزوية والعلمية",
    "📰 التحرير والإعلام الصحفي",
    "📄 فحص ومقارنة الوثائق",
    "🧭 القيادة والتخطيط والتقويم",
    "🌿 الاستشارات الحياتية والديكور",
    "📊 لوحة المؤشرات البيانية",
    "🗄️ الأرشيف الدائم والبحث"
])

# ==========================================
# 0. السكرتير الصوتي والشخصي
# ==========================================
with tab0:
    st.header("السكرتير الشخصي المباشر (صوت، ملفات، وصور)")
    audio_record = st.audio_input("🎙️ تسجيل صوتي مباشر:")
    sec_file = st.file_uploader("📎 أرفق ملف أو صورة متعلقة بالمهمة (اختياري):", type=["pdf", "docx", "png", "jpg", "jpeg"], key="sec_file")
    sec_text = st.text_area("أو اكتب رسالتك/توجيهك بالتفصيل:", placeholder="اكتب موضوعك، تعميم ترغب بصياغته، أو مسألة تريد ترتيبها...", height=100)
    
    if st.button("تنفيذ المهمة عبر السكرتير", type="primary"):
        if audio_record or sec_file or sec_text.strip():
            with st.spinner("السكرتير يعالج المدخلات ويصيغ الرد..."):
                try:
                    sys_inst = f"""أنت سكرتيري ومساعدي التنفيذي الخاص والشامل. أسلوبك: {secretary_mode}.
مهمتك إنجاز المطلوب بدقة إدارية ولغوية رفيعة.
تنبيه حازم: يُمنع منعاً باتاً كتابة أي توقيتات زمنية صوتية (مثل 01:23) في النص نهائياً."""
                    
                    contents = []
                    if audio_record:
                        audio_raw = audio_record.read()
                        contents.append(types.Part.from_bytes(data=audio_raw, mime_type="audio/wav"))
                        
                    payload = prepare_multimodal_payload(sys_inst, sec_text if sec_text.strip() else "نفذ المطلوب بناءً على الصوت أو الملف المرفق.", sec_file)
                    contents.extend(payload)
                    
                    res = generate_with_retry(contents)
                    reply = clean_text_output(res.text)
                    
                    title = sec_text[:30] if sec_text.strip() else (sec_file.name if sec_file else "مهمة صوتية")
                    save_record("محادثة شخصية", title, sec_text, reply)
                    
                    st.markdown("### 💬 رد السكرتير التنفيذي:")
                    st.markdown(reply)
                    if voice_output:
                        aud = text_to_audio_bytes(reply)
                        if aud:
                            st.audio(aud, format="audio/mp3")
                    docx_out = create_docx_download(reply, "مخرجات السكرتير")
                    st.download_button("📥 تحميل المخرجات (Word)", docx_out, file_name="Secretary_Output.docx")
                except Exception as err:
                    st.error(f"حدث خطأ: {err}")
        else:
            st.warning("يرجى إدخال صوت، كتابة نص، أو رفع ملف.")

# ==========================================
# 1. البحوث الحوزوية والعلمية
# ==========================================
with tab1:
    st.header("كتابة وتحقيق الأبحاث الحوزوية والمقالات العلمية")
    c1, c2 = st.columns(2)
    with c1:
        r_type = st.selectbox("المجال التخصصي:", ["بحث فقهي / أصولي استدلالي", "بحث كلامي وعقائدي", "دراسة قرآنية وحديثية", "تحقيق تراثي ورجالي", "مقال فكري وفلسفي", "بحث أكاديمي محكم"])
    with c2:
        r_meth = st.selectbox("المنهجية المعتمدة:", ["استدلالي حوزوي رصين (أقوال، أدلة، مناقشة، المختار)", "تحقيقي تراثي بالمصادر", "مقال فكري تحليلي"])
    
    topic = st.text_input("موضوع البحث أو القضية:")
    res_file = st.file_uploader("📎 أرفق مخطوطة، صورة صفحة، أو وثيقة علمية (PDF / Word / صورة):", type=["pdf", "docx", "png", "jpg", "jpeg"], key="res_file")
    r_notes = st.text_area("النصوص المقتبسة، الروايات، الأقوال، أو المحاور الخاصة:", height=100)
    
    if st.button("كتابة وتأصيل البحث"):
        if topic.strip() or res_file or r_notes.strip():
            with st.spinner("جاري التحقيق الاستدلالي الرصين..."):
                try:
                    sys_inst = f"""أنت باحث ومحقق حوزوي وأكاديمي خبير.
المجال: {r_type} | المنهج: {r_meth} | الموضوع: {topic}
المطلوب: تحرير محل النزاع، تفريع الأدلة، مناقشة الأقوال (إن قيل... قلنا)، واستخلاص الرأي المختار بدقة مع ثبت المصادر."""
                    payload = prepare_multimodal_payload(sys_inst, r_notes, res_file)
                    res = generate_with_retry(payload)
                    reply = clean_text_output(res.text)
                    
                    save_record("بحث حوزوي/علمي", topic if topic else "بحث علمي", r_notes, reply)
                    st.markdown("### 📜 النص العلمي المحرر:")
                    st.markdown(reply)
                    docx_res = create_docx_download(reply, f"بحث: {topic}")
                    st.download_button("📥 تحميل البحث (Word)", docx_res, file_name="Research.docx")
                except Exception as err:
                    st.error(f"خطأ: {err}")
        else:
            st.warning("يرجى كتابة الموضوع أو إرفاق ملف.")

# ==========================================
# 2. التحرير والإعلام الصحفي
# ==========================================
with tab2:
    st.header("صياغة الأخبار والبيانات الصحفية باحترافية")
    c_t, c_n = st.columns(2)
    with c_t:
        n_type = st.selectbox("القالب الإعلامي:", ["خبر صحفي (هرم مقلوب)", "منشور منصات تواصل (Facebook / X)", "بيان صحفي وتصريح رسمي", "تغطية إخبارية موسعة"])
    with c_n:
        n_tone = st.selectbox("نبرة التحرير:", ["احترافي رصين وجذاب", "حماسي وتفاعلي", "رسمي ومؤسسي دقيق"])
        
    press_file = st.file_uploader("📎 أرفق ملصق الفعالية، صورة الحدث، أو جدول الأعمال (اختياري):", type=["pdf", "docx", "png", "jpg", "jpeg"], key="press_file")
    n_facts = st.text_area("تفاصيل ووقائع الحدث أو الأرقام البارزة:", height=100)
    
    if st.button("صياغة المادة الإعلامية"):
        if n_facts.strip() or press_file:
            with st.spinner("جاري صياغة الخبر الصحفي..."):
                try:
                    sys_inst = f"""أنت رئيس تحرير وصحفي محترف. القالب: {n_type} | النبرة: {n_tone}.
المطلوب: 3 مقترحات عناوين جذابة، متن الخبر المتوازن، ونسخة مهيأة للسوشيال ميديا مع الوسوم المناسبة."""
                    payload = prepare_multimodal_payload(sys_inst, n_facts, press_file)
                    res = generate_with_retry(payload)
                    reply = clean_text_output(res.text)
                    
                    save_record("إعلام وصحافة", n_facts[:30] if n_facts else "خبر صحفي", n_facts, reply)
                    st.markdown("### 📰 المادة الصحفية الجاهزة:")
                    st.markdown(reply)
                    docx_res = create_docx_download(reply, "المادة الإعلامية")
                    st.download_button("📥 تحميل المادة الصحفية (Word)", docx_res, file_name="Press_Release.docx")
                except Exception as err:
                    st.error(f"خطأ: {err}")
        else:
            st.warning("يرجى إدخال تفاصيل الحدث أو رفع ملف.")

# ==========================================
# 3. فحص ومقارنة الوثائق
# ==========================================
with tab3:
    st.header("فحص وتحليل ومقارنة الوثائق والصور")
    doc_mode = st.radio("نوع العملية:", ["تدقيق وثيقة واحدة أو صورة", "مقارنة وثيقتين لكشف الفروقات والتعارضات"], horizontal=True)
    
    if doc_mode == "تدقيق وثيقة واحدة أو صورة":
        up_file = st.file_uploader("ارفع الوثيقة أو الصورة المراد تدقيقها:", type=["pdf", "docx", "png", "jpg", "jpeg"], key="doc_single")
        q_text = st.text_input("المطلوب استخراجه أو تدقيقه:", placeholder="مثال: اكتشف الثغرات، دقق لغوياً، لخص القرارات...")
        if st.button("بدء التدقيق"):
            if up_file and q_text:
                with st.spinner("جاري فحص الوثيقة..."):
                    try:
                        sys_inst = f"أنت خبير تدقيق إداري وقانوني. المطلوب: {q_text}."
                        payload = prepare_multimodal_payload(sys_inst, q_text, up_file)
                        res = generate_with_retry(payload)
                        reply = clean_text_output(res.text)
                        
                        save_record("فحص وثائق", up_file.name, q_text, reply)
                        st.markdown("### 📋 التقرير الصادر:")
                        st.markdown(reply)
                        docx_res = create_docx_download(reply, "تقرير فحص وثيقة")
                        st.download_button("📥 تحميل التقرير (Word)", docx_res, file_name="Doc_Report.docx")
                    except Exception as err:
                        st.error(f"خطأ: {err}")
            else:
                st.warning("يرجى رفع الملف وتحديد المطلوب.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            fa = st.file_uploader("الوثيقة الأصلية / المسودة الأولى:", type=["pdf", "docx"], key="fa")
        with col2:
            fb = st.file_uploader("الوثيقة المعدلة / المسودة الثانية:", type=["pdf", "docx"], key="fb")
        if st.button("مقارنة الوثيقتين"):
            if fa and fb:
                with st.spinner("جاري المقارنة واستخراج التغييرات..."):
                    try:
                        ta = extract_text_from_file(fa)
                        tb = extract_text_from_file(fb)
                        p = f"قارن بين الوثيقتين بدقة واستخرج جدول التعديلات، الثغرات، والتعارضات:\n[الوثيقة 1]:\n{ta[:5000]}\n\n[الوثيقة 2]:\n{tb[:5000]}"
                        res = generate_with_retry([p])
                        reply = clean_text_output(res.text)
                        
                        save_record("مقارنة وثائق", f"{fa.name} VS {fb.name}", "مقارنة نسختين", reply)
                        st.markdown("### 🔍 تقرير المقارنة:")
                        st.markdown(reply)
                        docx_res = create_docx_download(reply, "تقرير المقارنة")
                        st.download_button("📥 تحميل التقرير (Word)", docx_res, file_name="Comparison.docx")
                    except Exception as err:
                        st.error(f"خطأ: {err}")

# ==========================================
# 4. القيادة والتخطيط والتقويم
# ==========================================
with tab4:
    st.header("إدارة الأعمال، التخطيط الاستراتيجي، والتقييم والتقويم")
    mgmt_mode = st.radio("المهمة:", ["بناء خطة استراتيجية ومؤشرات أداء", "تقييم وتقويم الأداء ومعالجة الانحرافات", "حلول وتوجيه قيادي"], horizontal=True)
    
    mgmt_file = st.file_uploader("📎 أرفق ملف الخطة السابقة أو تقارير الإنجاز (اختياري):", type=["pdf", "docx", "png", "jpg", "jpeg"], key="mgmt_file")
    mgmt_notes = st.text_area("أدخل الأهداف أو بيانات المخطط مقابل المنجز أو التحدي الإداري:", height=100)
    
    if st.button("تنفيذ التحليل القيادي والإداري"):
        if mgmt_notes.strip() or mgmt_file:
            with st.spinner("جاري إعداد التحليل الإداري..."):
                try:
                    sys_inst = f"""أنت خبير قيادة وإدارة استراتيجية وعملياتية.
المجال المطلوب: {mgmt_mode}.
قدم مخرجات تنظيمية واضحة تشمل جداول مراحل، مصفوفات RACI، ومؤشرات قياس SMART."""
                    payload = prepare_multimodal_payload(sys_inst, mgmt_notes, mgmt_file)
                    res = generate_with_retry(payload)
                    reply = clean_text_output(res.text)
                    
                    save_record("إدارة وقيادة", mgmt_mode, mgmt_notes[:30], reply)
                    st.markdown("### 🧭 المخرج القيادي المعتمد:")
                    st.markdown(reply)
                    docx_res = create_docx_download(reply, "التقرير الإداري والقيادي")
                    st.download_button("📥 تحميل التقرير (Word)", docx_res, file_name="Management_Plan.docx")
                except Exception as err:
                    st.error(f"خطأ: {err}")
        else:
            st.warning("يرجى كتابة البيانات أو رفع ملف.")

# ==========================================
# 5. الاستشارات الحياتية والديكور
# ==========================================
with tab5:
    st.header("🌿 المستشار التخصصي: ديكور المنازل، الزراعة، الصحة، والعلاقات")
    consult_type = st.selectbox(
        "مجال الاستشارة:",
        [
            "🏡 تصميم وديكورات المنازل والمساحات",
            "🌱 استشارات زراعية ونباتات وأسمدة",
            "🩺 إرشادات ونمط حياة صحي عام",
            "🤝 علاقات اجتماعية وأسرية وبناء الذات"
        ]
    )
    consult_file = st.file_uploader("📎 أرفق صورة للمساحة/النبات أو تقرير (اختياري):", type=["pdf", "docx", "png", "jpg", "jpeg"], key="consult_file")
    consult_notes = st.text_area("تفاصيل السؤال، الأبعاد، أو الحالة المراد استشارتها:", height=100)
    
    if st.button("طلب الاستشارة التخصصية"):
        if consult_notes.strip() or consult_file:
            with st.spinner("المستشار التخصصي يحلل البيانات..."):
                try:
                    sys_inst = f"""أنت مستشار خبير في {consult_type}.
قدم تحليلاً عملياً ومباشراً:
- في الديكور: تحليل الإضاءة، الألوان، توزيع الأثاث، واختيار الخامات.
- في الزراعة: تشخيص الحالة، التربة، جدول الري، والتسميد المناسب.
- في الصحة: نصائح نمط الحياة مع التنويه بمراجعة المختص.
- في العلاقات: حلول حكيمة ومتزنة."""
                    payload = prepare_multimodal_payload(sys_inst, consult_notes, consult_file)
                    res = generate_with_retry(payload)
                    reply = clean_text_output(res.text)
                    
                    save_record(f"استشارة: {consult_type}", consult_notes[:30] if consult_notes else consult_type, consult_notes, reply)
                    st.markdown("### 💡 الرأي والاستشارة التخصصية:")
                    st.markdown(reply)
                    docx_res = create_docx_download(reply, f"استشارة - {consult_type}")
                    st.download_button("📥 تحميل الاستشارة (Word)", docx_res, file_name="Consultation.docx")
                except Exception as err:
                    st.error(f"خطأ: {err}")
        else:
            st.warning("يرجى كتابة السؤال أو إرفاق صورة/ملف.")

# ==========================================
# 6. اللوحة البيانية
# ==========================================
with tab6:
    st.header("📊 لوحة قياس الأداء والمتابعة البيانية")
    default_df = {
        "المسار / المهمة": ["الشؤون الشخصية والسكرتارية", "البحوث والتحقيق", "الإعلام والنشر", "التدقيق الإداري", "المشاريع الحياتية والتطوير"],
        "المستهدف (%)": [100, 100, 100, 100, 100],
        "المتحقق الفعلي (%)": [95, 90, 85, 92, 80]
    }
    df = st.data_editor(pd.DataFrame(default_df), num_rows="dynamic")
    if not df.empty:
        df["الفجوة (%)"] = df["المستهدف (%)"] - df["المتحقق الفعلي (%)"]
        fig = px.bar(
            df, 
            x="المسار / المهمة", 
            y=["المتحقق الفعلي (%)", "الفجوة (%)"],
            title="مقارنة الإنجاز الفعلي مقابل الفجوة المتبقية",
            barmode="stack",
            color_discrete_sequence=["#2ecc71", "#e74c3c"]
        )
        st.plotly_chart(fig, use_container_width=True)

# ==========================================
# 7. الأرشيف الدائم والبحث (SQLite)
# ==========================================
with tab_arch:
    st.header("🗄️ الأرشيف الدائم وقاعدة البيانات")
    st.write("استعرض، ابحث، أو أعد تحميل أي بحث، مستند، أو محادثة تم حفظها في النظام.")
    
    col_s1, col_s2 = st.columns([3, 1])
    with col_s1:
        s_query = st.text_input("🔍 ابحث في الأرشيف (بالكلمة أو العنوان أو المحتوى):", placeholder="اكتب للبحث...")
    with col_s2:
        cat_filter = st.selectbox(
            "تصفية حسب التصنيف:",
            ["الكل", "محادثة شخصية", "بحث حوزوي/علمي", "إعلام وصحافة", "فحص وثائق", "مقارنة وثائق", "إدارة وقيادة", "استشارة: 🏡 تصميم وديكورات المنازل والمساحات", "استشارة: 🌱 استشارات زراعية ونباتات وأسمدة", "استشارة: 🩺 إرشادات ونمط حياة صحي عام", "استشارة: 🤝 علاقات اجتماعية وأسرية وبناء الذات"]
        )
        
    records = get_records(s_query, cat_filter)
    st.caption(f"عدد السجلات المطابقة: {len(records)}")
    
    if records:
        for r in records:
            r_id, r_time, r_cat, r_title, r_prompt, r_response = r
            with st.expander(f"📌 [{r_cat}] {r_title} | 🕒 {r_time}"):
                st.markdown(f"**المدخلات:**\n{r_prompt}")
                st.markdown("---")
                st.markdown(f"**النتيجة:**\n{r_response}")
                
                col_d1, col_d2 = st.columns([2, 1])
                with col_d1:
                    d_file = create_docx_download(r_response, r_title)
                    st.download_button("📥 تحميل نسخة Word مجدداً", d_file, file_name=f"Archive_{r_id}.docx", key=f"dl_{r_id}")
                with col_d2:
                    if st.button("🗑️ حذف من الأرشيف", key=f"del_{r_id}"):
                        delete_record(r_id)
                        st.rerun()
    else:
        st.info("لا توجد سجلات محفوظة مطابقة لبحثك.")
